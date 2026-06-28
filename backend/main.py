from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from pdf2docx import Converter
import json, io, re, httpx, tempfile, os
from dotenv import load_dotenv
load_dotenv()
DEFAULT_API_KEY = os.getenv("OPENROUTER_KEY", "")

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL          = "google/gemini-2.5-flash-lite"


# ── 1. PDF → DOCX ─────────────────────────────────────────────────────────────
def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path  = os.path.join(tmp, "in.pdf")
        docx_path = os.path.join(tmp, "in.docx")
        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        with open(docx_path, "rb") as f:
            return f.read()


# ── 2. Extraire le texte pour Gemini ─────────────────────────────────────────
def extraire_texte(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for i, p in enumerate(doc.paragraphs):
        t = "".join(r.text for r in p.runs).strip()
        if t:
            lines.append(f"[§{i}] {t}")
    return "\n".join(lines)


# ── 3. Gemini retourne les remplacements ──────────────────────────────────────
async def get_remplacements(texte: str, infos: dict, api_key: str) -> dict:
    infos_str = "\n".join(f"- {k}: {v}" for k, v in infos.items() if str(v).strip())

    prompt = f"""Tu es un expert en documents légaux marocains.

Voici le texte du document (lignes numérotées) :
{texte}

Informations à insérer (mot pour mot) :
{infos_str}

MISSION : Pour chaque champ vide, retourne le numéro de paragraphe [§N] et le texte COMPLET et FINAL de ce paragraphe avec les valeurs insérées.

RÈGLES :
- Remplace les …………… par les vraies valeurs
- Garde TOUT le reste du texte identique
- Pour la ligne qui contient "MONTANT_EN_LETTRES" ou "(en toutes lettres)" → remplace par "Le montant du loyer initial est fixé à la somme de LOYER_LETTRES_PLACEHOLDER" — utilise exactement le texte LOYER_LETTRES_PLACEHOLDER sans accolades ni crochets
- Pour "(Ville)" → ville de signature  
- Pour "(Date)" → date de signature
- Pour "[Nombre]" → nombre d'exemplaires
- Pour "[Adresses des parties]" → adresses bailleur et locataire
- Pour "(Nature du tribunal/Ville)" → ville du tribunal
- Première ligne "M, Mme, Mlle…………………Né (e) le" → bailleur (inclure civilité + nom + date + lieu)
- Deuxième ligne "M, Mme, Mlle…………………Né (e) le" → locataire (inclure civilité + nom + date + lieu)
- Première ligne "Demeurant à……………" → adresse bailleur
- Deuxième ligne "Demeurant à……………" → adresse locataire

Réponds UNIQUEMENT avec un JSON valide :
{{"2": "texte complet final du paragraphe 2", "5": "texte complet final du paragraphe 5"}}"""

    async with httpx.AsyncClient(timeout=90) as client:
        resp = await client.post(
            OPENROUTER_URL,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={"model": MODEL, "messages": [{"role": "user", "content": prompt}],
                  "max_tokens": 4000, "temperature": 0}
        )

    if resp.status_code != 200:
        raise HTTPException(status_code=502, detail=f"Erreur API : {resp.text[:300]}")

    contenu = resp.json()["choices"][0]["message"]["content"].strip()
    contenu = re.sub(r"^```(?:json)?\s*|\s*```$", "", contenu.strip())
    try:
        result = json.loads(contenu)
        return result if isinstance(result, dict) else {}
    except Exception:
        m = re.search(r'\{.*\}', contenu, re.DOTALL)
        return json.loads(m.group(0)) if m else {}


# ── Constantes XML bordures ───────────────────────────────────────────────────
_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

NO_BORDER_TC = (
    f'<w:tcBorders xmlns:w="{_NS}">'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tcBorders>'
)

NO_BORDER_TBL = (
    f'<w:tblBorders xmlns:w="{_NS}">'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)


# ── 4. Modifier le DOCX ───────────────────────────────────────────────────────
def modifier_docx(docx_bytes: bytes, mapping: dict) -> bytes:
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document(io.BytesIO(docx_bytes))

    # ── Remplacer les paragraphes ──────────────────────────────────────────
    for idx_str, nouveau_texte in mapping.items():
        try:
            idx  = int(idx_str)
            para = doc.paragraphs[idx]
            runs = [r for r in para.runs if r.text.strip()]
            if not runs:
                continue

            ref_run   = runs[0]
            bold      = ref_run.bold
            italic    = ref_run.italic
            font_size = ref_run.font.size
            font_name = ref_run.font.name

            for r in para.runs:
                r.text = ""

            para.runs[0].text   = nouveau_texte
            para.runs[0].bold   = bold
            para.runs[0].italic = italic
            if font_size:
                para.runs[0].font.size = font_size
            if font_name:
                para.runs[0].font.name = font_name

            print(f"✅ §{idx} → {nouveau_texte[:60]}")

        except Exception as e:
            print(f"❌ §{idx_str} erreur : {e}")

    # ── Supprimer les bordures des tableaux ───────────────────────────────
    # FIX : assigner tblStyle=TableNormal (style sans bordures) ET nettoyer
    # tblLook pour éviter que Word applique TableGrid par défaut
    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn("w:tblPr"))

        # 1. Forcer le style TableNormal (zéro bordures) au lieu de TableGrid
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            tblPr.remove(tblStyle)
        new_style = etree.fromstring(
            f'<w:tblStyle xmlns:w="{_NS}" w:val="TableNormal"/>'
        )
        tblPr.insert(0, new_style)

        # 2. Désactiver tblLook (styles conditionnels firstRow/firstColumn)
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is not None:
            tblPr.remove(tblLook)
        tblPr.append(etree.fromstring(
            f'<w:tblLook xmlns:w="{_NS}" w:val="0000" '
            'w:firstRow="0" w:lastRow="0" w:firstColumn="0" '
            'w:lastColumn="0" w:noHBand="1" w:noVBand="1"/>'
        ))

        # 3. Remplacer tblBorders
        old_tbl_b = tblPr.find(qn("w:tblBorders"))
        if old_tbl_b is not None:
            tblPr.remove(old_tbl_b)
        tblPr.append(etree.fromstring(NO_BORDER_TBL))

        # 4. Remplacer tcBorders + supprimer shd dans chaque cellule
        for row in table.rows:
            for cell in row.cells:
                tc   = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = etree.SubElement(tc, qn("w:tcPr"))

                old_b = tcPr.find(qn("w:tcBorders"))
                if old_b is not None:
                    tcPr.remove(old_b)
                tcPr.append(etree.fromstring(NO_BORDER_TC))

                # Supprimer le shading (fond blanc forcé peut interagir avec le style)
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    tcPr.remove(shd)

    # ── Nettoyer les civilités devant les noms ─────────────────────────────
    # FIX : M\.?\s* matchait 'M' dans 'Monsieur' → utiliser M\. ou M suivi d'espace/virgule
    CIVILITES = re.compile(
        r"^(M,\s*Mme,\s*Mlle\s*"
        r"|Mme\.?[ \t]*"
        r"|Mlle\.?[ \t]*"
        r"|M\.[ \t]+"          # M. suivi d'espace
        r"|M,[ \t]*"           # M, (dans liste de civilités)
        r"|Monsieur[ \t]+"     # Monsieur suivi d'espace obligatoire
        r"|Madame[ \t]+"
        r"|Mademoiselle[ \t]+)",
        re.IGNORECASE
    )
    for para in doc.paragraphs:
        texte = "".join(r.text for r in para.runs)
        if "Né (e) le" in texte or "Né(e) le" in texte:
            nouveau = CIVILITES.sub("", texte)
            if nouveau != texte:
                para.runs[0].text = nouveau
                for r in para.runs[1:]:
                    r.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Conversion montant → toutes lettres ──────────────────────────────────────
def nombre_en_lettres(n: int) -> str:
    units = ["","un","deux","trois","quatre","cinq","six","sept","huit","neuf",
             "dix","onze","douze","treize","quatorze","quinze","seize",
             "dix-sept","dix-huit","dix-neuf"]
    tens  = ["","","vingt","trente","quarante","cinquante",
             "soixante","soixante","quatre-vingt","quatre-vingt"]
    if n == 0: return "zéro"
    r = ""
    if n >= 1000:
        m = n // 1000
        r += "mille" if m == 1 else nombre_en_lettres(m) + " mille"
        n %= 1000
        if n: r += " "
    if n >= 100:
        m = n // 100
        r += "cent" if m == 1 else units[m] + " cent"
        n %= 100
        if n: r += " "
    if n >= 20:
        t, u = n // 10, n % 10
        if t == 7:   r += "soixante-" + units[10 + u]
        elif t == 9: r += "quatre-vingt-" + units[10 + u]
        else:
            r += tens[t]
            if u == 1 and t != 8: r += "-et-un"
            elif u:               r += "-" + units[u]
    elif n > 0:
        r += units[n]
    return r.strip()

def loyer_en_lettres(val: str) -> str:
    try:
        n = int(re.sub(r"[^\d]", "", val))
        return nombre_en_lettres(n).upper() + " DIRHAMS"
    except Exception:
        return val.upper()


# ── Endpoint ──────────────────────────────────────────────────────────────────
@app.post("/remplir-pdf")
async def remplir_pdf(
    fichier:        UploadFile = File(...),
    infos:          str        = Form(...),
    openrouter_key: str        = Form(default=""),
    groq_key:       str        = Form(default=""),
):
    api_key = openrouter_key.strip() or groq_key.strip() or DEFAULT_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="Clé API OpenRouter manquante.")

    infos_dict = json.loads(infos)
    pdf_bytes  = await fichier.read()

    docx_bytes = pdf_to_docx(pdf_bytes)
    texte      = extraire_texte(docx_bytes)
    mapping    = await get_remplacements(texte, infos_dict, api_key)

    # ── Substituer le placeholder loyer — Gemini peut retourner plusieurs variantes
    loyer_val = infos_dict.get("Loyer mensuel (DH)", "")
    lettres   = loyer_en_lettres(loyer_val) if loyer_val else ""

    LOYER_PATTERNS = [
        "LOYER_LETTRES_PLACEHOLDER",
        "{LOYER_LETTRES}",
        "{{LOYER_LETTRES}}",
        "MONTANT_EN_LETTRES",
        "[MONTANT_EN_LETTRES]",
        "{MONTANT_EN_LETTRES}",
    ]
    for k in mapping:
        for pat in LOYER_PATTERNS:
            if pat in mapping[k]:
                mapping[k] = mapping[k].replace(pat, lettres)
                break

    print(f"\n📋 {len(mapping)} paragraphes à modifier : {list(mapping.keys())}\n")
    docx_rempli = modifier_docx(docx_bytes, mapping)

    return StreamingResponse(
        io.BytesIO(docx_rempli),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=document_rempli.docx"}
    )


@app.get("/")
def root():
    return {"status": "ok", "version": "v19"}